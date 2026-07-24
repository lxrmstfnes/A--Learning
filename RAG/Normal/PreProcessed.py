#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档预处理 (PreProcessed)
=========================

预处理流程:
    1. 提取 PDF / Word（.pdf / .docx / .doc）文本，记录页码（Word 为伪页）
    2. 清洗与规范化提取文本
    3. 递归字符分割（chunk_size=1000, overlap=200）
    4. 基于字符偏移，建立文本块与来源页码的映射

用法:
    python PreProcessed.py
    python PreProcessed.py --input data/某文件.pdf
    python PreProcessed.py --input data/某文件.docx
    python PreProcessed.py --input data/ --output processed/
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from pypdf import PdfReader


# =============================================================================
# 路径与默认参数
# =============================================================================

RAG_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT_DIR = RAG_ROOT / "data"
DEFAULT_OUTPUT_DIR = RAG_ROOT / "processed"

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
PAGE_SEPARATOR = "\n\n"

# 支持的文档扩展名
SUPPORTED_DOC_EXTENSIONS = {".pdf", ".doc", ".docx"}

# Word 无真实页码时，按约该字符数切成伪页，便于后续分批
WORD_PSEUDO_PAGE_CHARS = 1800

# 递归分割优先级：段落 -> 行 -> 中文句读 -> 英文句点 -> 空格 -> 字符
DEFAULT_SEPARATORS = ["\n\n", "\n", "。", "！", "？", "；", ". ", " ", ""]


# =============================================================================
# 数据结构
# =============================================================================


@dataclass
class PageRecord:
    """单页 PDF 提取结果。"""

    page_number: int
    text: str
    is_empty: bool = False
    error: Optional[str] = None


@dataclass
class PageSpan:
    """全文中某一页文本对应的字符区间（左闭右开）。"""

    page_number: int
    char_start: int
    char_end: int


@dataclass
class TextChunk:
    """分割后的文本块及页码映射。"""

    chunk_id: int
    text: str
    char_start: int
    char_end: int
    source_pages: List[int] = field(default_factory=list)


@dataclass
class PreprocessResult:
    """单份文档的完整预处理结果。"""

    source_file: str
    total_pages: int
    valid_pages: int
    empty_pages: List[int]
    error_pages: List[int]
    chunk_size: int
    chunk_overlap: int
    full_text_length: int
    pages: List[PageRecord]
    chunks: List[TextChunk]
    created_at: str = field(default_factory=lambda: datetime.now().isoformat(timespec="seconds"))


# =============================================================================
# 1. 文档文本提取（PDF / Word）
# =============================================================================


def normalize_page_text(raw: Optional[str]) -> str:
    """清洗单页文本：统一换行、去除首尾空白、压缩连续空行。"""
    if not raw:
        return ""

    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pages_from_pdf(pdf_path: Path) -> Tuple[List[PageRecord], int]:
    """
    逐页提取 PDF 文本并记录页码。

    返回:
        (页面记录列表, PDF 总页数)
    """
    reader = PdfReader(str(pdf_path))
    total_pages = len(reader.pages)
    records: List[PageRecord] = []

    for index in range(total_pages):
        page_number = index + 1
        try:
            raw_text = reader.pages[index].extract_text()
            text = normalize_page_text(raw_text)
            is_empty = len(text) == 0

            records.append(
                PageRecord(
                    page_number=page_number,
                    text=text,
                    is_empty=is_empty,
                )
            )
        except Exception as exc:  # noqa: BLE001 - 单页失败不中断整份文档
            records.append(
                PageRecord(
                    page_number=page_number,
                    text="",
                    is_empty=True,
                    error=str(exc),
                )
            )

    return records, total_pages


def _paragraphs_from_docx(doc_path: Path) -> List[str]:
    """用 python-docx 提取段落与表格文本。"""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "处理 .docx 需要安装 python-docx，请执行: pip install python-docx"
        ) from exc

    document = Document(str(doc_path))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    return parts


def _run_command(command: Sequence[str], timeout: int = 120) -> subprocess.CompletedProcess[str]:
    """运行外部命令，捕获 stdout/stderr。"""
    return subprocess.run(
        list(command),
        check=False,
        capture_output=True,
        text=True,
        timeout=timeout,
    )


def _extract_doc_via_textutil(doc_path: Path) -> Optional[str]:
    """macOS textutil：.doc → txt。"""
    if not shutil.which("textutil"):
        return None

    with tempfile.TemporaryDirectory(prefix="rag_doc_") as tmp_dir:
        out_path = Path(tmp_dir) / f"{doc_path.stem}.txt"
        result = _run_command(
            ["textutil", "-convert", "txt", str(doc_path), "-output", str(out_path)]
        )
        if result.returncode != 0 or not out_path.exists():
            return None
        return out_path.read_text(encoding="utf-8", errors="ignore")


def _extract_doc_via_libreoffice(doc_path: Path) -> Optional[str]:
    """LibreOffice / soffice：.doc → txt。"""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None

    with tempfile.TemporaryDirectory(prefix="rag_doc_") as tmp_dir:
        result = _run_command(
            [
                soffice,
                "--headless",
                "--norestore",
                "--convert-to",
                "txt:Text",
                "--outdir",
                tmp_dir,
                str(doc_path),
            ],
            timeout=180,
        )
        if result.returncode != 0:
            return None

        candidates = list(Path(tmp_dir).glob("*.txt"))
        if not candidates:
            return None
        return candidates[0].read_text(encoding="utf-8", errors="ignore")


def extract_docx_text(doc_path: Path) -> str:
    """提取 .docx 全文。"""
    return "\n\n".join(_paragraphs_from_docx(doc_path))


def extract_doc_text(doc_path: Path) -> str:
    """
    提取旧版 .doc 全文。

    优先顺序: macOS textutil → LibreOffice → 尝试按 docx 打开。
    """
    text = _extract_doc_via_textutil(doc_path)
    if text and text.strip():
        return text

    text = _extract_doc_via_libreoffice(doc_path)
    if text and text.strip():
        return text

    # 少数文件扩展名为 .doc 实为 OOXML
    try:
        text = extract_docx_text(doc_path)
        if text.strip():
            return text
    except Exception:  # noqa: BLE001
        pass

    raise RuntimeError(
        f"无法解析 .doc 文件: {doc_path.name}。"
        "请安装 LibreOffice（soffice），或在 macOS 上确保可用 textutil；"
        "也可先将该文件另存为 .docx / .pdf 后再入库。"
    )


def extract_word_text(word_path: Path) -> str:
    """按扩展名提取 Word 全文。"""
    suffix = word_path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(word_path)
    if suffix == ".doc":
        return extract_doc_text(word_path)
    raise ValueError(f"不支持的 Word 扩展名: {word_path}")


def split_text_into_pseudo_pages(
    text: str,
    max_chars: int = WORD_PSEUDO_PAGE_CHARS,
) -> List[PageRecord]:
    """
    将无页码文档按段落切成伪页，便于沿用「按页分批」逻辑。

    优先在段落边界断开；单段超长时再硬切。
    """
    text = normalize_page_text(text)
    if not text:
        return [PageRecord(page_number=1, text="", is_empty=True)]

    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    if not paragraphs:
        paragraphs = [text]

    pages: List[PageRecord] = []
    buffer = ""

    def flush() -> None:
        nonlocal buffer
        content = normalize_page_text(buffer)
        if not content:
            buffer = ""
            return
        pages.append(
            PageRecord(
                page_number=len(pages) + 1,
                text=content,
                is_empty=False,
            )
        )
        buffer = ""

    for paragraph in paragraphs:
        candidate = f"{buffer}\n\n{paragraph}".strip() if buffer else paragraph
        if buffer and len(candidate) > max_chars:
            flush()
            candidate = paragraph

        if len(candidate) <= max_chars:
            buffer = candidate
            continue

        # 单段过长：按 max_chars 硬切
        flush()
        start = 0
        while start < len(paragraph):
            piece = paragraph[start : start + max_chars]
            pages.append(
                PageRecord(
                    page_number=len(pages) + 1,
                    text=piece,
                    is_empty=False,
                )
            )
            start += max_chars

    flush()
    return pages or [PageRecord(page_number=1, text="", is_empty=True)]


def extract_pages_from_word(word_path: Path) -> Tuple[List[PageRecord], int]:
    """提取 Word 文本并切成伪页。"""
    raw = extract_word_text(word_path)
    pages = split_text_into_pseudo_pages(raw)
    return pages, len(pages)


def extract_pages_from_document(doc_path: Path) -> Tuple[List[PageRecord], int]:
    """按文件类型提取页面记录（PDF 真实页 / Word 伪页）。"""
    suffix = doc_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pages_from_pdf(doc_path)
    if suffix in {".doc", ".docx"}:
        return extract_pages_from_word(doc_path)
    raise ValueError(
        f"不支持的文件类型: {doc_path.name}（支持: {', '.join(sorted(SUPPORTED_DOC_EXTENSIONS))}）"
    )


# =============================================================================
# 2. 合并有效页文本
# =============================================================================


def build_full_text(pages: Sequence[PageRecord]) -> Tuple[str, List[PageSpan], List[int], List[int]]:
    """
    将非空页拼接为全文，并记录每页字符区间。

    返回:
        (full_text, page_spans, empty_pages, error_pages)
    """
    full_text = ""
    page_spans: List[PageSpan] = []
    empty_pages: List[int] = []
    error_pages: List[int] = []

    for page in pages:
        if page.error:
            error_pages.append(page.page_number)
        if page.is_empty:
            empty_pages.append(page.page_number)
            continue

        if full_text:
            full_text += PAGE_SEPARATOR

        char_start = len(full_text)
        full_text += page.text
        page_spans.append(
            PageSpan(
                page_number=page.page_number,
                char_start=char_start,
                char_end=len(full_text),
            )
        )

    return full_text, page_spans, empty_pages, error_pages


# =============================================================================
# 3. 递归字符分割器
# =============================================================================


class RecursiveCharacterTextSplitter:
    """
    递归字符文本分割器。

    优先在语义边界（段落、行、句读符）处切分，尽量保持 chunk_size 以内。
    """

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
        separators: Optional[Sequence[str]] = None,
    ) -> None:
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap 必须小于 chunk_size")

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = list(separators or DEFAULT_SEPARATORS)

    def split_text(self, text: str) -> List[str]:
        """将文本分割为多个 chunk 字符串。"""
        if not text:
            return []

        splits = self._split_text(text, self.separators)
        return self._merge_splits(splits)

    def split_text_with_offsets(self, text: str) -> List[Tuple[str, int, int]]:
        """分割文本并返回 (chunk_text, char_start, char_end)。"""
        chunks = self.split_text(text)
        if not chunks:
            return []

        result: List[Tuple[str, int, int]] = []
        search_from = 0

        for chunk in chunks:
            start = text.find(chunk, search_from)
            if start < 0:
                start = text.find(chunk)
            if start < 0:
                raise RuntimeError("无法在原文中定位 chunk 偏移，请检查分割逻辑。")

            end = start + len(chunk)
            result.append((chunk, start, end))
            search_from = max(start + 1, end - self.chunk_overlap)

        return result

    def _split_text(self, text: str, separators: Sequence[str]) -> List[str]:
        """递归按分隔符优先级切分过长片段。"""
        final_chunks: List[str] = []
        separator = separators[-1]
        next_separators: List[str] = []

        for index, candidate in enumerate(separators):
            if candidate == "":
                separator = candidate
                break
            if candidate in text:
                separator = candidate
                next_separators = list(separators[index + 1 :])
                break

        if separator:
            splits = text.split(separator)
        else:
            splits = list(text)

        merged: List[str] = []
        for index, split in enumerate(splits):
            if split:
                merged.append(split)
            if separator and index < len(splits) - 1:
                merged.append(separator)

        good_splits: List[str] = []
        for split in merged:
            if not split:
                continue
            if len(split) <= self.chunk_size:
                good_splits.append(split)
            else:
                if not next_separators:
                    good_splits.extend(self._hard_split(split))
                else:
                    good_splits.extend(self._split_text(split, next_separators))

        return good_splits

    def _hard_split(self, text: str) -> List[str]:
        """无合适分隔符时按固定长度硬切。"""
        return [
            text[index : index + self.chunk_size]
            for index in range(0, len(text), self.chunk_size)
        ]

    def _merge_splits(self, splits: Sequence[str]) -> List[str]:
        """将细粒度 split 合并为带 overlap 的 chunk。"""
        chunks: List[str] = []
        current: List[str] = []
        current_len = 0

        for split in splits:
            split_len = len(split)
            if current and current_len + split_len > self.chunk_size:
                chunk_text = "".join(current)
                if chunk_text:
                    chunks.append(chunk_text)

                while current and current_len > self.chunk_overlap:
                    removed = current.pop(0)
                    current_len -= len(removed)

                while current and current_len + split_len > self.chunk_size:
                    removed = current.pop(0)
                    current_len -= len(removed)

            current.append(split)
            current_len += split_len

        if current:
            chunks.append("".join(current))

        return chunks


# =============================================================================
# 4. 文本块与页码映射
# =============================================================================


def map_chunk_to_pages(char_start: int, char_end: int, page_spans: Sequence[PageSpan]) -> List[int]:
    """
    根据字符区间 [char_start, char_end) 匹配来源页码。

    任一 page span 与 chunk 区间有交集，即视为来源页。
    """
    pages: List[int] = []
    for span in page_spans:
        if span.char_start < char_end and span.char_end > char_start:
            pages.append(span.page_number)
    return pages


def split_and_map_pages(
    full_text: str,
    page_spans: Sequence[PageSpan],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[TextChunk]:
    """分割全文并为每个 chunk 建立页码映射。"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    split_results = splitter.split_text_with_offsets(full_text)

    chunks: List[TextChunk] = []
    for chunk_id, (text, char_start, char_end) in enumerate(split_results):
        chunks.append(
            TextChunk(
                chunk_id=chunk_id,
                text=text,
                char_start=char_start,
                char_end=char_end,
                source_pages=map_chunk_to_pages(char_start, char_end, page_spans),
            )
        )
    return chunks


# =============================================================================
# 预处理主流程
# =============================================================================


def preprocess_document(
    doc_path: Path,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> PreprocessResult:
    """执行单份文档（PDF / Word）的完整预处理流程。"""
    pages, total_pages = extract_pages_from_document(doc_path)
    full_text, page_spans, empty_pages, error_pages = build_full_text(pages)

    chunks = split_and_map_pages(
        full_text=full_text,
        page_spans=page_spans,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    return PreprocessResult(
        source_file=str(doc_path.resolve()),
        total_pages=total_pages,
        valid_pages=len(page_spans),
        empty_pages=empty_pages,
        error_pages=error_pages,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        full_text_length=len(full_text),
        pages=pages,
        chunks=chunks,
    )


def preprocess_pdf(
    pdf_path: Path,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> PreprocessResult:
    """兼容旧接口：等价于 preprocess_document。"""
    return preprocess_document(pdf_path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)


def save_preprocess_result(result: PreprocessResult, output_path: Path) -> None:
    """将预处理结果保存为 JSON。"""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    payload = {
        "source_file": result.source_file,
        "total_pages": result.total_pages,
        "valid_pages": result.valid_pages,
        "empty_pages": result.empty_pages,
        "error_pages": result.error_pages,
        "chunk_size": result.chunk_size,
        "chunk_overlap": result.chunk_overlap,
        "full_text_length": result.full_text_length,
        "chunk_count": len(result.chunks),
        "created_at": result.created_at,
        "pages": [asdict(page) for page in result.pages],
        "chunks": [asdict(chunk) for chunk in result.chunks],
    }
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _is_supported_document(path: Path) -> bool:
    """判断是否为可入库文档（排除 Office 临时锁文件）。"""
    name = path.name
    if name.startswith("~$") or name.startswith(".~"):
        return False
    return path.suffix.lower() in SUPPORTED_DOC_EXTENSIONS


def iter_document_files(input_path: Path) -> Iterable[Path]:
    """解析输入路径，返回待处理的 PDF / Word 文件列表。"""
    if input_path.is_file():
        if not _is_supported_document(input_path):
            raise ValueError(
                f"仅支持 {', '.join(sorted(SUPPORTED_DOC_EXTENSIONS))} 文件: {input_path}"
            )
        return [input_path]

    if not input_path.exists():
        raise FileNotFoundError(f"输入路径不存在: {input_path}")

    files: List[Path] = []
    for suffix in sorted(SUPPORTED_DOC_EXTENSIONS):
        files.extend(input_path.glob(f"*{suffix}"))
    # 去重并按文件名排序（macOS 可能对大小写扩展名各匹配一次）
    unique = sorted(
        {path.resolve(): path for path in files if _is_supported_document(path)}.values(),
        key=lambda p: p.name.lower(),
    )
    if not unique:
        raise FileNotFoundError(
            f"目录中未找到支持的文档（{', '.join(sorted(SUPPORTED_DOC_EXTENSIONS))}）: {input_path}"
        )
    return unique


def iter_pdf_files(input_path: Path) -> Iterable[Path]:
    """兼容旧接口：返回 PDF / Word 文档列表。"""
    return iter_document_files(input_path)


def default_output_file(doc_path: Path, output_dir: Path) -> Path:
    """根据文档文件名生成输出 JSON 路径。"""
    return output_dir / f"{doc_path.stem}.preprocessed.json"


def print_summary(result: PreprocessResult, output_path: Path) -> None:
    """打印单份文档预处理摘要。"""
    print("-" * 60)
    print(f"文件: {Path(result.source_file).name}")
    print(f"总页数: {result.total_pages} | 有效页: {result.valid_pages}")
    if result.empty_pages:
        print(f"空白页: {result.empty_pages}")
    if result.error_pages:
        print(f"异常页: {result.error_pages}")
    print(f"全文长度: {result.full_text_length} 字符")
    print(f"文本块数: {len(result.chunks)} (size={result.chunk_size}, overlap={result.chunk_overlap})")

    if result.chunks:
        sample = result.chunks[0]
        preview = sample.text.replace("\n", " ")[:80]
        print(f"示例块 #0 | 页码 {sample.source_pages} | {preview}...")

    print(f"输出: {output_path}")


# =============================================================================
# 命令行入口
# =============================================================================


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="文档预处理 — PDF/Word 提取 + 递归字符分割 + 页码映射"
    )
    parser.add_argument(
        "--input",
        type=str,
        default=str(DEFAULT_INPUT_DIR),
        help=f"PDF/Word 文件或目录（默认: {DEFAULT_INPUT_DIR}）",
    )
    parser.add_argument(
        "--output",
        type=str,
        default=str(DEFAULT_OUTPUT_DIR),
        help=f"预处理结果输出目录（默认: {DEFAULT_OUTPUT_DIR}）",
    )
    parser.add_argument("--chunk-size", type=int, default=CHUNK_SIZE, help="文本块大小")
    parser.add_argument("--chunk-overlap", type=int, default=CHUNK_OVERLAP, help="文本块重叠长度")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output).expanduser().resolve()

    try:
        doc_files = list(iter_document_files(input_path))
    except (FileNotFoundError, ValueError) as exc:
        print(f"[错误] {exc}", file=sys.stderr)
        sys.exit(1)

    print("=" * 60)
    print("  文档预处理 (PreProcessed) — PDF / Word")
    print("=" * 60)
    print(f"  输入: {input_path}")
    print(f"  输出目录: {output_dir}")
    print(f"  支持格式: {', '.join(sorted(SUPPORTED_DOC_EXTENSIONS))}")
    print(f"  分割参数: chunk_size={args.chunk_size}, overlap={args.chunk_overlap}")
    print("=" * 60)

    for doc_path in doc_files:
        try:
            result = preprocess_document(
                doc_path=doc_path,
                chunk_size=args.chunk_size,
                chunk_overlap=args.chunk_overlap,
            )
            output_path = default_output_file(doc_path, output_dir)
            save_preprocess_result(result, output_path)
            print_summary(result, output_path)
        except Exception as exc:  # noqa: BLE001
            print(f"[错误] 处理失败 {doc_path.name}: {exc}", file=sys.stderr)

    print("\n[完成] 预处理结束。")


if __name__ == "__main__":
    main()
